"""
Market Analysis Export Router — MARKET-EXPORT-001
Endpoints: /pdf, /docx, /xlsx, /md for market analysis reports.
Based on TZ v3.0 export requirements.
"""
import io
import json
import logging
from typing import Optional
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.db.session import get_db
from app.api.v1.routers.auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/uz-market/reports", tags=["market-analysis-export"])


def _get_report(report_id: str, db: Session, user) -> dict:
    """Fetch report from DB or raise 404."""
    from app.db.models.reports import ReportInstance
    if report_id == "latest":
        instance = (
            db.query(ReportInstance)
            .filter(ReportInstance.user_id == user.id)
            .order_by(ReportInstance.created_at.desc())
            .first()
        )
    else:
        instance = (
            db.query(ReportInstance)
            .filter(ReportInstance.id == int(report_id), ReportInstance.user_id == user.id)
            .first()
        )
    if not instance:
        raise HTTPException(status_code=404, detail="Report not found")
    return json.loads(instance.content) if isinstance(instance.content, str) else instance.content


def _sections_to_text(report: dict) -> str:
    """Convert report sections to plain text."""
    lines = []
    lines.append(f"MARKET ANALYSIS REPORT")
    lines.append(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}")
    lines.append("")
    rec = report.get("recommendation", "N/A")
    score = report.get("confidence_score", 0)
    lines.append(f"Recommendation: {rec.upper()} (Confidence: {score:.1f}%)")
    lines.append("")
    if report.get("executive_summary"):
        lines.append("EXECUTIVE SUMMARY")
        lines.append(report["executive_summary"])
        lines.append("")
    for sec in report.get("sections", []):
        lines.append(f"{sec.get('number', '')}. {sec.get('title', '')}")
        lines.append(sec.get("content", ""))
        lines.append("")
    return "\n".join(lines)


@router.get("/{report_id}/md", summary="Export report as Markdown")
async def export_markdown(report_id: str, db: Session = Depends(get_db), user=Depends(get_current_user)):
    report = _get_report(report_id, db, user)
    md_lines = [f"# Market Analysis Report\n"]
    md_lines.append(f"> Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}\n")
    rec = report.get("recommendation", "N/A")
    score = report.get("confidence_score", 0)
    md_lines.append(f"**Recommendation:** {rec.upper()} | **Confidence:** {score:.1f}%\n")
    if report.get("executive_summary"):
        md_lines.append(f"## Executive Summary\n\n{report['executive_summary']}\n")
    for sec in report.get("sections", []):
        md_lines.append(f"## {sec.get('number', '')}. {sec.get('title', '')}\n\n{sec.get('content', '')}\n")
    content = "\n".join(md_lines)
    buf = io.BytesIO(content.encode("utf-8"))
    return StreamingResponse(buf, media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="market-report-{report_id}.md"'})


@router.get("/{report_id}/pdf", summary="Export report as PDF")
async def export_pdf(report_id: str, db: Session = Depends(get_db), user=Depends(get_current_user)):
    report = _get_report(report_id, db, user)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_LEFT

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=25*mm, rightMargin=25*mm, topMargin=30*mm, bottomMargin=25*mm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("CustomTitle", parent=styles["Title"], fontSize=18, spaceAfter=12)
        h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, spaceAfter=8, spaceBefore=16)
        body_style = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10, spaceAfter=6, leading=14)
        story = []
        story.append(Paragraph("AI Capital — Market Analysis Report", title_style))
        story.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}", body_style))
        rec = report.get("recommendation", "N/A")
        score = report.get("confidence_score", 0)
        story.append(Paragraph(f"<b>Recommendation:</b> {rec.upper()} | <b>Confidence:</b> {score:.1f}%", body_style))
        story.append(Spacer(1, 12))
        if report.get("executive_summary"):
            story.append(Paragraph("Executive Summary", h2_style))
            story.append(Paragraph(report["executive_summary"].replace("\n", "<br/>"), body_style))
        for sec in report.get("sections", []):
            story.append(Paragraph(f"{sec.get('number','')}. {sec.get('title','')}", h2_style))
            content_safe = sec.get("content", "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
            story.append(Paragraph(content_safe, body_style))
        story.append(Spacer(1, 20))
        story.append(Paragraph("© AI Capital Management — автоматический отчёт", body_style))
        doc.build(story)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="market-report-{report_id}.pdf"'})
    except ImportError:
        raise HTTPException(status_code=501, detail="reportlab not installed")


@router.get("/{report_id}/docx", summary="Export report as DOCX")
async def export_docx(report_id: str, db: Session = Depends(get_db), user=Depends(get_current_user)):
    report = _get_report(report_id, db, user)
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()
        doc.add_heading("AI Capital — Market Analysis Report", level=0)
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}")
        rec = report.get("recommendation", "N/A")
        score = report.get("confidence_score", 0)
        p = doc.add_paragraph()
        p.add_run(f"Recommendation: {rec.upper()}").bold = True
        p.add_run(f" | Confidence: {score:.1f}%")
        if report.get("executive_summary"):
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(report["executive_summary"])
        for sec in report.get("sections", []):
            doc.add_heading(f"{sec.get('number','')}. {sec.get('title','')}", level=1)
            doc.add_paragraph(sec.get("content", ""))
        doc.add_paragraph("© AI Capital Management", style="Intense Quote")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="market-report-{report_id}.docx"'})
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx not installed")


@router.get("/{report_id}/xlsx", summary="Export report as XLSX")
async def export_xlsx(report_id: str, db: Session = Depends(get_db), user=Depends(get_current_user)):
    report = _get_report(report_id, db, user)
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment

        wb = Workbook()
        ws = wb.active
        ws.title = "Market Report"
        ws.append(["AI Capital — Market Analysis Report"])
        ws.merge_cells("A1:C1")
        ws["A1"].font = Font(bold=True, size=14)
        ws.append([])
        ws.append(["Recommendation", report.get("recommendation", "").upper()])
        ws.append(["Confidence Score", f"{report.get('confidence_score', 0):.1f}%"])
        ws.append(["Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M")])
        ws.append([])

        if report.get("macro_context"):
            ws.append(["— Macro Context —"])
            for k, v in report["macro_context"].items():
                ws.append([k, str(v)])
            ws.append([])

        ws.append(["Section #", "Title", "Content"])
        ws["A" + str(ws.max_row)].font = Font(bold=True)
        ws["B" + str(ws.max_row)].font = Font(bold=True)
        ws["C" + str(ws.max_row)].font = Font(bold=True)
        for sec in report.get("sections", []):
            ws.append([sec.get("number", ""), sec.get("title", ""), sec.get("content", "")])

        ws.column_dimensions["A"].width = 12
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 80

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="market-report-{report_id}.xlsx"'})
    except ImportError:
        raise HTTPException(status_code=501, detail="openpyxl not installed")
