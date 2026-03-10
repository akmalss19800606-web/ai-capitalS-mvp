"""
Сервис анализа DD-документов — извлечение данных из загруженных файлов.

Фаза 3, DD-002:
  - Анализ PDF (pdfplumber), DOCX (python-docx), XLSX (openpyxl)
  - Извлечение: даты, суммы, ИНН, стороны контракта
  - Типы: financial_report, charter, license, contract
  - Структурированный результат с risk_flags
"""

import io
import logging
import re
import uuid
from datetime import datetime, timezone
from typing import Optional

logger = logging.getLogger(__name__)

# ── Попытка импорта библиотек ──────────────────────────────────

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


# ── Шаблоны типов документов ───────────────────────────────────

DOCUMENT_TEMPLATES = {
    "financial_report": {
        "label": "Финансовый отчёт",
        "description": "Бухгалтерская отчётность, баланс, P&L",
        "extract_fields": ["revenue", "net_income", "total_assets", "total_liabilities", "period"],
    },
    "charter": {
        "label": "Устав компании",
        "description": "Учредительный документ юридического лица",
        "extract_fields": ["company_name", "inn", "charter_fund", "founders", "activities"],
    },
    "license": {
        "label": "Лицензия / разрешение",
        "description": "Государственная лицензия на вид деятельности",
        "extract_fields": ["license_number", "issue_date", "expiry_date", "issuing_authority", "activities"],
    },
    "contract": {
        "label": "Договор / контракт",
        "description": "Хозяйственный договор между сторонами",
        "extract_fields": ["parties", "contract_date", "contract_amount", "subject", "term"],
    },
}


class DocumentAnalysisService:
    """Сервис анализа DD-документов."""

    # Хранилище анализов (для MVP — in-memory)
    _analyses: dict[str, dict] = {}

    @classmethod
    async def analyze_document(
        cls,
        file_bytes: bytes,
        filename: str,
        doc_type: str = "auto",
    ) -> dict:
        """
        Анализ загруженного документа.

        Args:
            file_bytes: Содержимое файла.
            filename: Имя файла.
            doc_type: Тип документа (auto для автоопределения).

        Returns:
            Структурированный результат анализа.
        """
        doc_id = str(uuid.uuid4())[:8]
        extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        # Извлекаем текст из файла
        text = cls._extract_text(file_bytes, extension)

        # Автоопределение типа документа
        if doc_type == "auto":
            doc_type = cls._detect_doc_type(text, filename)

        # Извлекаем структурированные данные
        extracted = cls._extract_fields(text, doc_type)

        # Обнаруживаем риски
        risk_flags = cls._detect_risks(text, doc_type, extracted)

        # Формируем результат
        analysis = {
            "doc_id": doc_id,
            "filename": filename,
            "doc_type": doc_type,
            "doc_type_label": DOCUMENT_TEMPLATES.get(doc_type, {}).get("label", doc_type),
            "file_size_bytes": len(file_bytes),
            "text_length": len(text),
            "extracted_fields": extracted,
            "risk_flags": risk_flags,
            "summary": cls._generate_summary(extracted, doc_type),
            "analyzed_at": datetime.now(timezone.utc).isoformat(),
        }

        # Сохраняем в кэш
        cls._analyses[doc_id] = analysis

        return analysis

    @classmethod
    def get_analysis(cls, doc_id: str) -> Optional[dict]:
        """Получить результат анализа по ID."""
        return cls._analyses.get(doc_id)

    @classmethod
    def get_templates(cls) -> dict:
        """Получить список доступных шаблонов анализа."""
        return DOCUMENT_TEMPLATES

    # ── Извлечение текста ────────────────────────────────────

    @classmethod
    def _extract_text(cls, file_bytes: bytes, extension: str) -> str:
        """Извлечь текст из файла."""
        if extension == "pdf" and PDFPLUMBER_AVAILABLE:
            return cls._extract_pdf(file_bytes)
        elif extension in ("docx", "doc") and DOCX_AVAILABLE:
            return cls._extract_docx(file_bytes)
        elif extension in ("xlsx", "xls") and OPENPYXL_AVAILABLE:
            return cls._extract_xlsx(file_bytes)
        elif extension == "txt":
            return file_bytes.decode("utf-8", errors="replace")
        else:
            # Попробуем как текст
            try:
                return file_bytes.decode("utf-8", errors="replace")
            except Exception:
                return ""

    @classmethod
    def _extract_pdf(cls, file_bytes: bytes) -> str:
        """Извлечь текст из PDF через pdfplumber."""
        text_parts = []
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages[:50]:  # Ограничение на 50 страниц
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            logger.error(f"Ошибка извлечения текста из PDF: {e}")
        return "\n".join(text_parts)

    @classmethod
    def _extract_docx(cls, file_bytes: bytes) -> str:
        """Извлечь текст из DOCX."""
        text_parts = []
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Извлекаем также из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
        except Exception as e:
            logger.error(f"Ошибка извлечения текста из DOCX: {e}")
        return "\n".join(text_parts)

    @classmethod
    def _extract_xlsx(cls, file_bytes: bytes) -> str:
        """Извлечь текст из XLSX."""
        text_parts = []
        try:
            wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            for sheet in wb.sheetnames[:5]:  # Первые 5 листов
                ws = wb[sheet]
                text_parts.append(f"=== Лист: {sheet} ===")
                for row in ws.iter_rows(max_row=200, values_only=True):
                    row_text = " | ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        text_parts.append(row_text)
        except Exception as e:
            logger.error(f"Ошибка извлечения текста из XLSX: {e}")
        return "\n".join(text_parts)

    # ── Определение типа документа ───────────────────────────

    @classmethod
    def _detect_doc_type(cls, text: str, filename: str) -> str:
        """Автоопределение типа документа по содержимому."""
        text_lower = text.lower()
        fname_lower = filename.lower()

        # По ключевым словам
        if any(kw in text_lower for kw in ["баланс", "выручка", "прибыль", "отчёт", "доход", "расход"]):
            return "financial_report"
        if any(kw in text_lower for kw in ["устав", "учредител", "уставный фонд", "основной вид деятельности"]):
            return "charter"
        if any(kw in text_lower for kw in ["лицензия", "разрешение", "лицензиат", "выдана"]):
            return "license"
        if any(kw in text_lower for kw in ["договор", "контракт", "стороны", "заказчик", "исполнитель"]):
            return "contract"

        # По имени файла
        if any(kw in fname_lower for kw in ["баланс", "finance", "financial", "отчет", "pnl"]):
            return "financial_report"
        if any(kw in fname_lower for kw in ["устав", "charter"]):
            return "charter"
        if any(kw in fname_lower for kw in ["лицензия", "license"]):
            return "license"
        if any(kw in fname_lower for kw in ["договор", "contract"]):
            return "contract"

        return "financial_report"  # default

    # ── Извлечение полей ─────────────────────────────────────

    @classmethod
    def _extract_fields(cls, text: str, doc_type: str) -> dict:
        """Извлечь структурированные поля из текста."""
        fields = {}

        # Общие поля
        fields["dates"] = cls._extract_dates(text)
        fields["amounts"] = cls._extract_amounts(text)
        fields["inns"] = cls._extract_inns(text)

        # По типу документа
        if doc_type == "financial_report":
            fields.update(cls._extract_financial(text))
        elif doc_type == "charter":
            fields.update(cls._extract_charter(text))
        elif doc_type == "license":
            fields.update(cls._extract_license(text))
        elif doc_type == "contract":
            fields.update(cls._extract_contract(text))

        return fields

    @classmethod
    def _extract_dates(cls, text: str) -> list[str]:
        """Извлечь все даты из текста."""
        patterns = [
            r"\d{2}\.\d{2}\.\d{4}",  # DD.MM.YYYY
            r"\d{4}-\d{2}-\d{2}",     # YYYY-MM-DD
            r"\d{2}/\d{2}/\d{4}",     # DD/MM/YYYY
        ]
        dates = set()
        for pattern in patterns:
            dates.update(re.findall(pattern, text))
        return sorted(dates)[:20]

    @classmethod
    def _extract_amounts(cls, text: str) -> list[dict]:
        """Извлечь денежные суммы."""
        amounts = []
        # Суммы в сумах/долларах
        patterns = [
            (r"([\d\s,.]+)\s*(сум|сўм|UZS)", "UZS"),
            (r"\$\s*([\d\s,.]+)", "USD"),
            (r"([\d\s,.]+)\s*(долл|USD)", "USD"),
            (r"([\d\s,.]+)\s*(руб|RUB)", "RUB"),
        ]
        for pattern, currency in patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                amount_str = match.group(1).strip().replace(" ", "").replace(",", ".")
                try:
                    amount_val = float(amount_str)
                    if amount_val > 0:
                        amounts.append({"value": amount_val, "currency": currency, "raw": match.group(0).strip()})
                except ValueError:
                    continue
        return amounts[:20]

    @classmethod
    def _extract_inns(cls, text: str) -> list[str]:
        """Извлечь все ИНН (9 цифр) из текста."""
        return list(set(re.findall(r"\b(\d{9})\b", text)))[:10]

    @classmethod
    def _extract_financial(cls, text: str) -> dict:
        """Извлечь финансовые показатели."""
        result = {}
        patterns = {
            "revenue": [r"(?:выручка|revenue|оборот)[\s:]*?([\d\s,.]+)", ],
            "net_income": [r"(?:чистая прибыль|net income|прибыль)[\s:]*?([\d\s,.]+)"],
            "total_assets": [r"(?:итого активы|total assets|активы)[\s:]*?([\d\s,.]+)"],
        }
        for field, pats in patterns.items():
            for pat in pats:
                match = re.search(pat, text, re.IGNORECASE)
                if match:
                    result[field] = match.group(1).strip()
                    break
        return result

    @classmethod
    def _extract_charter(cls, text: str) -> dict:
        """Извлечь данные из устава."""
        result = {}
        # Наименование
        match = re.search(r"(?:наименование|название)[\s:]*(.+?)(?:\n|$)", text, re.IGNORECASE)
        if match:
            result["company_name"] = match.group(1).strip()
        # Уставный фонд
        match = re.search(r"уставн[а-я]*\s*фонд[\s:]*(.+?)(?:\n|$)", text, re.IGNORECASE)
        if match:
            result["charter_fund"] = match.group(1).strip()
        return result

    @classmethod
    def _extract_license(cls, text: str) -> dict:
        """Извлечь данные из лицензии."""
        result = {}
        match = re.search(r"(?:лицензия|разрешение)\s*(?:№|#)\s*(\S+)", text, re.IGNORECASE)
        if match:
            result["license_number"] = match.group(1).strip()
        match = re.search(r"(?:выдан[а-я]*|дата выдачи)[\s:]*(\d{2}[./]\d{2}[./]\d{4})", text, re.IGNORECASE)
        if match:
            result["issue_date"] = match.group(1)
        match = re.search(r"(?:действ[а-я]*\s*до|срок)[\s:]*(\d{2}[./]\d{2}[./]\d{4})", text, re.IGNORECASE)
        if match:
            result["expiry_date"] = match.group(1)
        return result

    @classmethod
    def _extract_contract(cls, text: str) -> dict:
        """Извлечь данные из договора."""
        result = {}
        # Стороны
        parties = []
        for match in re.finditer(r"(?:ООО|АО|ЧП|ИП)\s*[«\"](.*?)[»\"]", text):
            parties.append(match.group(0).strip())
        if parties:
            result["parties"] = parties[:4]
        # Дата договора
        match = re.search(r"(?:от|дата)\s*(\d{2}[./]\d{2}[./]\d{4})", text)
        if match:
            result["contract_date"] = match.group(1)
        # Сумма
        match = re.search(r"(?:сумма|стоимость|цена)\s*(?:договора|контракта)?[\s:]*?([\d\s,.]+)", text, re.IGNORECASE)
        if match:
            result["contract_amount"] = match.group(1).strip()
        return result

    # ── Обнаружение рисков ────────────────────────────────────

    @classmethod
    def _detect_risks(cls, text: str, doc_type: str, extracted: dict) -> list[str]:
        """Выявление рисковых индикаторов в документе."""
        flags = []
        text_lower = text.lower()

        # Общие риски
        if len(text) < 100:
            flags.append("ДОКУМЕНТ СЛИШКОМ КОРОТКИЙ — возможно, не удалось извлечь текст")

        if not extracted.get("dates"):
            flags.append("ДАТЫ НЕ НАЙДЕНЫ — невозможно определить актуальность")

        if not extracted.get("inns"):
            flags.append("ИНН НЕ НАЙДЕН — невозможно идентифицировать стороны")

        # Ключевые слова-индикаторы риска
        risk_keywords = {
            "ликвидация": "Упоминание ликвидации",
            "банкротство": "Упоминание банкротства",
            "арбитраж": "Упоминание арбитражного разбирательства",
            "штраф": "Упоминание штрафных санкций",
            "задолженность": "Упоминание задолженности",
            "просроч": "Упоминание просроченных обязательств",
        }
        for keyword, label in risk_keywords.items():
            if keyword in text_lower:
                flags.append(label)

        return flags

    # ── Генерация сводки ─────────────────────────────────────

    @classmethod
    def _generate_summary(cls, extracted: dict, doc_type: str) -> str:
        """Генерация текстовой сводки анализа."""
        parts = []
        template = DOCUMENT_TEMPLATES.get(doc_type, {})
        parts.append(f"Тип документа: {template.get('label', doc_type)}")

        if extracted.get("dates"):
            parts.append(f"Найдено дат: {len(extracted['dates'])} (первая: {extracted['dates'][0]})")

        if extracted.get("amounts"):
            total = len(extracted["amounts"])
            parts.append(f"Найдено денежных сумм: {total}")

        if extracted.get("inns"):
            parts.append(f"Найдено ИНН: {', '.join(extracted['inns'])}")

        if extracted.get("parties"):
            parts.append(f"Стороны: {', '.join(extracted['parties'])}")

        return ". ".join(parts) if parts else "Анализ завершён, извлечённых данных мало."
